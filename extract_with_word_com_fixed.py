#!/usr/bin/env python3

import win32com.client
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
import os

# -----------------------------------
# WORD CONSTANTS (define explicitly)
# -----------------------------------
# Page numbering constants
wdRestartContinuous = 0
wdRestartPage = 2
wdRestartSection = 1

# Page number style constants
wdPageNumberStyleLowercaseRoman = 14
wdPageNumberStyleUppercaseRoman = 13

# Alignment constants
wdAlignParagraphCenter = 1

# Field type constants
wdFieldPage = 33

# Break type constants
wdSectionBreakNextPage = 2
wdPageBreak = 7

# Header/Footer constants
wdHeaderFooterPrimary = 1


# -----------------------------------
# CLEAN TEXT (REMOVE ALL LINE BREAKS AND CLEAN)
# -----------------------------------
def clean_text(text):
    if not text:
        return ""
    # Remove control characters and clean up
    cleaned = "".join(
        c for c in text
        if c >= " " or c in ("\n", "\t")
    ).strip()

    # Remove trailing slashes and extra whitespace
    cleaned = cleaned.rstrip("/\\").strip()

    return cleaned


def clean_title_text(text):
    """
    Special cleaning for titles - removes ALL line breaks and normalizes spaces
    """
    if not text:
        return ""

    # Remove ALL types of line breaks and normalize spaces
    cleaned = text.replace("\r\n", " ").replace("\r", " ").replace("\n", " ")

    # Replace multiple spaces with single space
    import re
    cleaned = re.sub(r'\s+', ' ', cleaned)

    # Remove control characters except spaces
    cleaned = "".join(c for c in cleaned if c >= " ")

    # Remove trailing slashes and extra whitespace
    cleaned = cleaned.rstrip("/\\").strip()

    return cleaned


# -----------------------------------
# WORD APP
# -----------------------------------
def get_word():
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    return word


# -----------------------------------
# EXTRACT HEADINGS (COLLECT ALL CONSECUTIVE HEADING 2 TEXT)
# -----------------------------------
def extract_headings(doc_path):
    print("\n📖 Extracting headings with consecutive Heading 2 collection")

    word = get_word()
    doc = word.Documents.Open(os.path.abspath(doc_path))
    doc.Repaginate()

    headings = []
    total = doc.Paragraphs.Count

    i = 1
    while i <= total:
        try:
            para = doc.Paragraphs(i)
            text = clean_text(para.Range.Text)
            style = para.Style.NameLocal.lower()

            # Look for Heading 1 (Chapter numbers)
            if "heading 1" in style and text:
                print(f"\n   📍 Found Heading 1: '{text}'")

                # ---- Chapter with Heading 2 title(s) ----
                if text.lower().startswith("chapter"):
                    print(f"   🔍 Processing chapter: '{text}'")

                    chapter = text
                    chapter_title_parts = []
                    j = i + 1

                    # Look for ALL consecutive Heading 2 paragraphs
                    print(f"   🔍 Looking for consecutive Heading 2 titles starting from paragraph {j}...")

                    while j <= total:
                        try:
                            p = doc.Paragraphs(j)
                            p_style = p.Style.NameLocal.lower()

                            # Get RAW text from the entire paragraph range
                            raw_text = p.Range.Text
                            cleaned_text = clean_title_text(raw_text)

                            print(f"   📝 Paragraph {j}: '{cleaned_text}' (style: {p_style})")

                            if not raw_text or not raw_text.strip():
                                print(f"   ⏭️ Empty paragraph {j}, skipping")
                                j += 1
                                continue

                            # Found Heading 2 - collect this part of the title
                            if "heading 2" in p_style:
                                chapter_title_parts.append(cleaned_text)
                                print(f"   ✅ Collected Heading 2 part: '{cleaned_text}'")
                                j += 1
                                continue  # Keep looking for more Heading 2 parts

                            # If we hit another Heading 1, stop looking
                            elif "heading 1" in p_style:
                                print(f"   🛑 Hit another Heading 1, stopping search")
                                break

                            # If we hit any other style, stop looking for more title parts
                            else:
                                print(f"   🛑 Hit non-heading style, stopping title collection")
                                break

                        except Exception as e:
                            print(f"   ⚠️ Error processing paragraph {j}: {e}")
                            break

                    # Build the complete title from all parts
                    if chapter_title_parts:
                        # Join all title parts with a space
                        complete_title = " ".join(chapter_title_parts)
                        full = f"{chapter}: {complete_title}"
                        print(f"   🎉 Complete chapter title: '{full}'")
                    else:
                        full = chapter
                        print(f"   ⚠️ No Heading 2 found, using chapter only: '{full}'")

                    # Final cleaning to ensure single line
                    full = clean_title_text(full)
                    page = para.Range.Information(3)

                    headings.append({"text": full, "page": page})
                    print(f"   📌 Added to TOC: '{full}' → page {page}")

                    i = j
                    continue

                # ---- Single-line heading (Discography, Books, etc.) ----
                else:
                    page = para.Range.Information(3)
                    # Clean single-line headings too
                    clean_heading = clean_title_text(text)
                    headings.append({"text": clean_heading, "page": page})
                    print(f"   📌 Single-line heading: '{clean_heading}' → page {page}")

            # Also capture standalone Heading 2 (if not part of a chapter)
            elif "heading 2" in style and text:
                page = para.Range.Information(3)
                # Clean standalone Heading 2
                clean_heading = clean_title_text(para.Range.Text)
                headings.append({"text": clean_heading, "page": page})
                print(f"   📌 Standalone Heading 2: '{clean_heading}' → page {page}")

        except Exception as e:
            print(f"   ⚠️ Error processing paragraph {i}: {e}")

        i += 1

    doc.Close(False)
    word.Quit()

    print(f"\n   ✅ Total headings extracted: {len(headings)}")
    for idx, h in enumerate(headings, 1):
        print(f"   {idx}. {h['text']} → page {h['page']}")

    return headings


# -----------------------------------
# BUILD TOC (GUARANTEED SINGLE LINE)
# -----------------------------------
def build_toc_doc(headings, toc_path):
    print("\n📝 Building TOC with guaranteed single-line titles")

    doc = Document()

    # Title
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title.add_run("CONTENTS")
    run.font.name = "Georgia"
    run.font.size = Pt(20)
    run.bold = True

    doc.add_paragraph()

    for h in headings:
        para = doc.add_paragraph()

        # Hanging indent
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)

        # Tab stop with dots
        para.paragraph_format.tab_stops.add_tab_stop(
            Inches(6),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS
        )

        # Get the title text and ensure it's single line
        title_text = h["text"]

        # Final safety check - remove any remaining line breaks
        title_text = title_text.replace("\r\n", " ").replace("\r", " ").replace("\n", " ")
        title_text = " ".join(title_text.split())  # Normalize all whitespace

        print(f"   📝 TOC entry: '{title_text}' → {h['page']}")

        # Add title text (guaranteed single line)
        run = para.add_run(title_text)
        run.font.name = "Georgia"
        run.font.size = Pt(12)

        # Add tab and page number
        page_run = para.add_run(f"\t{h['page']}")
        page_run.bold = True

    doc.save(toc_path)
    print("   ✅ TOC created with guaranteed single-line entries")

    return toc_path


# -----------------------------------
# APPLY ROMAN PAGINATION (FIXED FOR COPYRIGHT PAGE START)
# -----------------------------------
def apply_roman_pagination(doc):
    print("\n🔢 Applying Roman pagination starting from copyright page")

    try:
        sections = doc.Sections
        print(f"   📊 Total sections: {sections.Count}")

        if sections.Count < 2:
            print("   ⚠️ Not enough sections")
            return

        # Section 1 (Title) → NO numbering
        print("   🔧 Section 1 (Title): NO page numbering")
        sec1 = sections.Item(1)
        try:
            # Clear headers and footers for title page
            sec1.Headers.Item(wdHeaderFooterPrimary).Range.Delete()
            sec1.Footers.Item(wdHeaderFooterPrimary).Range.Delete()

            # Ensure no page numbering on title section
            sec1.PageSetup.DifferentFirstPageHeaderFooter = True

            print("   ✅ Section 1: No numbering applied")

        except Exception as e:
            print(f"   ⚠️ Section 1 setup error: {e}")

        # Section 2 → Roman numerals starting from "i" on copyright page
        print("   🔧 Section 2 (Copyright + TOC): Roman numerals starting from 'i'")
        sec2 = sections.Item(2)

        # CRITICAL: Unlink from previous section
        try:
            sec2.Headers.Item(wdHeaderFooterPrimary).LinkToPrevious = False
            sec2.Footers.Item(wdHeaderFooterPrimary).LinkToPrevious = False
            print("   ✅ Section 2: Unlinked from previous section")
        except Exception as e:
            print(f"   ⚠️ Section 2 unlinking failed: {e}")

        # Clear existing footer content
        footer = sec2.Footers.Item(wdHeaderFooterPrimary)
        footer.Range.Delete()

        # Method 1: PageNumbers collection with proper restart
        try:
            print("   🔧 Method 1: PageNumbers collection...")

            # CRITICAL: Set page setup properties BEFORE adding page numbers
            page_setup = sec2.PageSetup

            # Force restart numbering at 1 (which becomes Roman "i")
            try:
                page_setup.RestartPageNumbering = True
                page_setup.PageNumberStart = 1
                print("   ✅ Page restart: True, Start: 1")
            except Exception as setup_err:
                print(f"   ⚠️ Page setup properties failed: {setup_err}")
                # Try alternative approach
                try:
                    page_setup.PageNumberingType = wdRestartPage
                    print("   ✅ Alternative restart method applied")
                except:
                    print("   ⚠️ All restart methods failed")

            # Add page numbers with Roman format
            page_nums = footer.PageNumbers
            page_nums.Add(PageNumberAlignment=wdAlignParagraphCenter)
            page_nums.NumberStyle = wdPageNumberStyleLowercaseRoman
            page_nums.StartingNumber = 1  # Copyright page = "i"

            print("   ✅ Roman page numbers added (i, ii, iii...)")

            # Force field updates
            doc.Fields.Update()
            doc.Repaginate()

            print("   ✅ Method 1 successful")

        except Exception as e1:
            print(f"   ⚠️ Method 1 failed: {e1}")

            # Method 2: Manual field with restart
            try:
                print("   🔧 Method 2: Manual field insertion...")
                footer.Range.Delete()

                # Set restart properties
                try:
                    sec2.PageSetup.RestartPageNumbering = True
                    sec2.PageSetup.PageNumberStart = 1
                except:
                    pass

                # Insert Roman numeral field
                footer_range = footer.Range
                field = footer_range.Fields.Add(
                    Range=footer_range,
                    Type=wdFieldPage,
                    PreserveFormatting=False
                )
                field.Code.Text = "PAGE \\* ROMAN \\* LOWER"
                field.Update()

                # Center the page number
                footer.Range.ParagraphFormat.Alignment = wdAlignParagraphCenter

                # Update document
                doc.Fields.Update()
                doc.Repaginate()

                print("   ✅ Method 2 successful")

            except Exception as e2:
                print(f"   ⚠️ Method 2 failed: {e2}")

                # Method 3: Force restart with field formula
                try:
                    print("   🔧 Method 3: Force restart formula...")
                    footer.Range.Delete()

                    # Insert a field that forces restart from 1
                    footer.Range.Text = "{ PAGE \\* ROMAN \\* LOWER }"
                    footer.Range.Fields.Update()
                    footer.Range.ParagraphFormat.Alignment = wdAlignParagraphCenter

                    # Try to force restart
                    try:
                        sec2.PageSetup.PageNumberStart = 1
                        sec2.PageSetup.RestartPageNumbering = True
                    except:
                        pass

                    doc.Fields.Update()
                    doc.ActiveWindow.View.ShowFieldCodes = False
                    doc.Repaginate()

                    print("   ✅ Method 3 applied")

                except Exception as e3:
                    print(f"   ❌ All methods failed: {e3}")

        # Final verification and updates
        try:
            print("   🔄 Final verification...")

            # Ensure field codes are hidden
            doc.ActiveWindow.View.ShowFieldCodes = False

            # Multiple updates to ensure restart takes effect
            doc.Repaginate()
            doc.Fields.Update()
            doc.Repaginate()

            print("   ✅ Final updates completed")

        except Exception as e:
            print(f"   ⚠️ Final update error: {e}")

    except Exception as e:
        print(f"   ❌ Roman pagination error: {e}")


# -----------------------------------
# ASSEMBLE FINAL DOC (IMPROVED)
# -----------------------------------
def assemble_final(title_doc, copyright_doc, toc_doc, output):
    print("\n📚 Assembling final document with single-line titles")

    word = get_word()
    doc = word.Documents.Add()

    try:
        # Title page (Section 1) - NO page numbering
        if os.path.exists(title_doc):
            print("   📄 Inserting title page (Section 1)...")
            r = doc.Content
            r.Collapse(0)
            r.InsertFile(os.path.abspath(title_doc))

            # Insert section break (creates Section 2 for copyright + TOC)
            r = doc.Content
            r.Collapse(0)
            r.InsertBreak(wdSectionBreakNextPage)
            print("   ✅ Title page + section break inserted")

        # Copyright page (Section 2, page "i")
        if os.path.exists(copyright_doc):
            print("   📄 Inserting copyright page (will be Roman 'i')...")
            r = doc.Content
            r.Collapse(0)
            r.InsertFile(os.path.abspath(copyright_doc))

            # Page break before TOC (stays in same section)
            r = doc.Content
            r.Collapse(0)
            r.InsertBreak(wdPageBreak)
            print("   ✅ Copyright page inserted")

        # TOC (Section 2, pages "ii", "iii", etc.)
        print("   📄 Inserting TOC (will be Roman 'ii', 'iii'...)...")
        r = doc.Content
        r.Collapse(0)
        r.InsertFile(os.path.abspath(toc_doc))
        print("   ✅ TOC inserted")

        # Apply Roman pagination starting from copyright page
        apply_roman_pagination(doc)

        # Final processing
        print("   🔄 Final document processing...")
        try:
            doc.ActiveWindow.View.ShowFieldCodes = False
            doc.Repaginate()
            doc.Fields.Update()
            doc.Repaginate()
        except Exception as e:
            print(f"   ⚠️ Final processing warning: {e}")

        # Save
        print("   💾 Saving document...")
        doc.SaveAs2(os.path.abspath(output))
        doc.Close(False)
        word.Quit()

        print("\n🎉 DOCUMENT COMPLETED!")
        print(f"📁 File: {output}")
        print("\n📋 Expected result:")
        print("   • Title page: No number")
        print("   • Copyright page: Roman 'i'")
        print("   • TOC pages: Roman 'ii', 'iii', 'iv'...")
        print("   • ALL chapter titles on single lines")

    except Exception as e:
        try:
            word.Quit()
        except:
            pass
        print(f"❌ Assembly error: {e}")


# -----------------------------------
# MAIN
# -----------------------------------
if __name__ == "__main__":
    book_file = "HITS AND HAPPINESS FINAL 2 Format MOM Discog.docx"
    title_file = "HH Title.docx"
    copyright_file = "HH Copyright page.docx"
    toc_temp = "temp_toc.docx"
    output_file = "HITS AND HAPPINESS FINAL 2 Format MOM Discog TOC.docx"

    headings = extract_headings(book_file)

    if headings:
        build_toc_doc(headings, toc_temp)
        assemble_final(title_file, copyright_file, toc_temp, output_file)

        if os.path.exists(toc_temp):
            os.remove(toc_temp)
