#!/usr/bin/env python3
"""
Alignment Comparison Demo
Shows the difference between old (incorrect) and new (correct) alignment methods
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER

def create_comparison_demo():
    """Create a demo showing old vs new alignment methods"""
    
    doc = Document()
    
    # Set up page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ALIGNMENT COMPARISON DEMO")
    title_run.font.name = 'Georgia'
    title_run.font.size = Pt(16)
    title_run.bold = True
    
    doc.add_paragraph()
    
    # OLD METHOD (INCORRECT)
    old_header = doc.add_paragraph()
    old_header_run = old_header.add_run("❌ OLD METHOD (Incorrect - Manual Dots)")
    old_header_run.font.name = 'Georgia'
    old_header_run.font.size = Pt(14)
    old_header_run.bold = True
    old_header_run.font.color.rgb = None  # Red would be nice but keeping it simple
    
    # Example with old method
    old_para1 = doc.add_paragraph()
    old_run1 = old_para1.add_run("Chapter 1: Introduction")
    old_run1.font.name = 'Georgia'
    old_run1.font.size = Pt(12)
    
    # Manual dots (this is what was wrong)
    dots = "." * 50
    old_dots1 = old_para1.add_run(dots + "5")
    old_dots1.font.name = 'Georgia'
    old_dots1.font.size = Pt(12)
    
    old_para2 = doc.add_paragraph()
    old_run2 = old_para2.add_run("Chapter 2: A Very Long Chapter Title That Causes Alignment Issues")
    old_run2.font.name = 'Georgia'
    old_run2.font.size = Pt(12)
    
    # Manual dots (notice how this doesn't align properly)
    dots2 = "." * 20
    old_dots2 = old_para2.add_run(dots2 + "17")
    old_dots2.font.name = 'Georgia'
    old_dots2.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # NEW METHOD (CORRECT)
    new_header = doc.add_paragraph()
    new_header_run = new_header.add_run("✅ NEW METHOD (Correct - Tab Stops)")
    new_header_run.font.name = 'Georgia'
    new_header_run.font.size = Pt(14)
    new_header_run.bold = True
    
    # Example with new method
    new_para1 = doc.add_paragraph()
    
    # Set up tab stop for right-aligned page number with dot leaders
    tab_stops1 = new_para1.paragraph_format.tab_stops
    tab_stops1.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    
    new_run1 = new_para1.add_run("Chapter 1: Introduction")
    new_run1.font.name = 'Georgia'
    new_run1.font.size = Pt(12)
    
    # Tab and page number (properly right-aligned)
    new_page1 = new_para1.add_run("\t5")
    new_page1.font.name = 'Georgia'
    new_page1.font.size = Pt(12)
    new_page1.bold = True
    
    new_para2 = doc.add_paragraph()
    
    # Set up tab stop for right-aligned page number with dot leaders
    tab_stops2 = new_para2.paragraph_format.tab_stops
    tab_stops2.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    
    new_run2 = new_para2.add_run("Chapter 2: A Very Long Chapter Title That Causes Alignment Issues")
    new_run2.font.name = 'Georgia'
    new_run2.font.size = Pt(12)
    
    # Tab and page number (notice how this DOES align properly)
    new_page2 = new_para2.add_run("\t17")
    new_page2.font.name = 'Georgia'
    new_page2.font.size = Pt(12)
    new_page2.bold = True
    
    doc.add_paragraph()
    
    # Explanation
    explanation = doc.add_paragraph()
    explanation_run = explanation.add_run("EXPLANATION:")
    explanation_run.font.name = 'Georgia'
    explanation_run.font.size = Pt(12)
    explanation_run.bold = True
    
    doc.add_paragraph("• Old method: Manually calculates dots based on title length")
    doc.add_paragraph("• Problem: Different title lengths = misaligned page numbers")
    doc.add_paragraph("• New method: Uses Word's built-in tab stops")
    doc.add_paragraph("• Result: All page numbers align perfectly to the right margin")
    doc.add_paragraph("• Professional: Automatic dot leaders, consistent spacing")
    
    doc.save("alignment_comparison_demo.docx")
    print("✓ Created 'alignment_comparison_demo.docx' to show the difference")

if __name__ == "__main__":
    create_comparison_demo()