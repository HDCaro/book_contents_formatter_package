import re
import json
import requests
from collections import defaultdict
from docx import Document
from docx.shared import Inches, Pt

# ---------------- CONFIG ---------------- #

DOCX_INPUT = "Hits And Happiness Final 2 Discog.docx"
DOCX_OUTPUT = "Hits And Happiness Final 2 Discog index.docx"

CACHE_FILE = "discogs_cache.json"

WORDS_PER_PAGE = 600
MIN_OCCURRENCES = 2

DISCOGS_TOKEN = ""  # optional
VERBOSE = True

# ---------------- LOGGING ---------------- #

def log(msg):
    if VERBOSE:
        print(msg)

# ---------------- CACHE ---------------- #

def load_cache():
    try:
        with open(CACHE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except:
        return {}

def save_cache(cache):
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, indent=2)

cache = load_cache()

# ---------------- TEXT ---------------- #

def extract_text(docx):
    doc = Document(docx)
    return "\n".join(p.text for p in doc.paragraphs)

def split_pages(text):
    words = text.split()
    pages = [
        " ".join(words[i:i + WORDS_PER_PAGE])
        for i in range(0, len(words), WORDS_PER_PAGE)
    ]
    log(f"[INFO] Total pages estimated: {len(pages)}")
    return pages

# ---------------- NORMALIZATION ---------------- #

def normalize_title(text):
    if "," in text:
        parts = [p.strip() for p in text.split(",")]
        parts.reverse()
        fixed = " ".join(parts)
        log(f"[NORMALIZE] '{text}' → '{fixed}'")
        return fixed
    return text

def clean_query(text):
    cleaned = re.sub(r"[^\w\s]", "", text)
    return cleaned

# ---------------- DISCOGS ---------------- #

def discogs_search(query):
    original_query = query
    query = normalize_title(query)
    query = clean_query(query)

    if query in cache:
        log(f"[CACHE] '{query}' → {cache[query]}")
        return cache[query]

    headers = {"User-Agent": "IndexBuilder/1.0"}
    if DISCOGS_TOKEN:
        headers["Authorization"] = f"Discogs token={DISCOGS_TOKEN}"

    queries = [
        query,
        f"{query} song",
        f"{query} track"
    ]

    for q in queries:
        log(f"[DISCOGS QUERY] {q}")

        try:
            r = requests.get(
                "https://api.discogs.com/database/search",
                params={"q": q},
                headers=headers,
                timeout=5
            )

            data = r.json()
            results = data.get("results", [])

            log(f"[DISCOGS RESULT] {len(results)} results")

            if len(results) > 0:
                log(f"[MATCH FOUND] '{original_query}' → '{query}'")
                cache[query] = True
                return True

        except Exception as e:
            log(f"[ERROR] Discogs request failed: {e}")

    log(f"[NO MATCH] '{original_query}'")
    cache[query] = False
    return False

# ---------------- DETECTION ---------------- #

def extract_candidates(text):
    pattern = r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\b"
    matches = re.findall(pattern, text)
    log(f"[CANDIDATES] Found {len(matches)} candidates")
    return matches

def classify_entries(text):
    candidates = extract_candidates(text)
    results = []

    for c in candidates:
        log(f"\n[CHECK] Candidate: {c}")

        parts = c.split()

        if len(parts) < 2:
            log("  → Skipped (too short)")
            continue

        # Try WORK
        normalized = normalize_title(c)

        if discogs_search(normalized):
            log(f"  → Classified as WORK: {normalized}")
            results.append(("WORK", normalized))
            continue

        # Otherwise PERSON
        if len(parts) <= 3:
            entry = f"{parts[-1]}, {' '.join(parts[:-1])}"
            log(f"  → Classified as PERSON: {entry}")
            results.append(("PERSON", entry))
        else:
            log("  → Skipped (too long for person)")

    return results

# ---------------- INDEX ---------------- #

def build_index(pages):
    index = defaultdict(set)

    for i, page in enumerate(pages, 1):
        log(f"\n================ PAGE {i} ================")

        entries = classify_entries(page)

        for typ, entry in entries:
            index[entry].add(i)
            log(f"[INDEX ADD] {entry} → page {i}")

    # filter
    filtered = {
        k: v for k, v in index.items()
        if len(v) >= MIN_OCCURRENCES
    }

    log(f"\n[INFO] Entries before filter: {len(index)}")
    log(f"[INFO] Entries after filter: {len(filtered)}")

    return filtered

# ---------------- FORMAT ---------------- #

def compress(pages):
    pages = sorted(pages)
    ranges = []
    start = prev = pages[0]

    for p in pages[1:]:
        if p == prev + 1:
            prev = p
        else:
            ranges.append((start, prev))
            start = prev = p

    ranges.append((start, prev))

    return ", ".join(
        f"{a}–{b}" if a != b else str(a)
        for a, b in ranges
    )

def split_sections(index):
    people, works = [], []

    for k in sorted(index):
        line = f"{k}, {compress(list(index[k]))}"

        if "," in k:
            people.append(line)
        else:
            works.append(line)

    log(f"[INFO] People entries: {len(people)}")
    log(f"[INFO] Works entries: {len(works)}")

    return people, works

# ---------------- DOCX ---------------- #

def add_header(doc, text, size):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Georgia"
    r.font.size = Pt(size)

def add_lines(doc, lines):
    for l in lines:
        p = doc.add_paragraph(l)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.space_after = Pt(0)

def create_doc(sections, out):
    doc = Document()

    title = doc.add_paragraph()
    r = title.add_run("INDEX")
    r.bold = True
    r.font.name = "Georgia"
    r.font.size = Pt(20)
    title.alignment = 1

    people, works = sections

    add_header(doc, "People", 16)
    add_lines(doc, people)

    add_header(doc, "Works", 16)
    add_lines(doc, works)

    doc.save(out)

# ---------------- MAIN ---------------- #

def main():
    print("\n=== START INDEX GENERATION ===\n")

    text = extract_text(DOCX_INPUT)
    pages = split_pages(text)

    index = build_index(pages)

    sections = split_sections(index)

    create_doc(sections, DOCX_OUTPUT)

    save_cache(cache)

    print("\n=== DONE ===")

if __name__ == "__main__":
    main()