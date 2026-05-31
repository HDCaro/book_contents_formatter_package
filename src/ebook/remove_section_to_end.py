#!/usr/bin/env python3

from docx import Document
import sys
import os


def delete_from_heading_to_end(input_file, heading_text, output_file):
    """
    Remove the specified heading and everything after it.
    """

    doc = Document(input_file)

    start_element = None

    for p in doc.paragraphs:
        if p.text.strip().lower() == heading_text.strip().lower():
            start_element = p._element
            break

    if start_element is None:
        raise ValueError(
            f'Heading "{heading_text}" was not found.'
        )

    body = doc.element.body

    delete_mode = False
    elements_to_remove = []

    for elem in body:
        if elem == start_element:
            delete_mode = True

        if delete_mode:
            elements_to_remove.append(elem)

    for elem in elements_to_remove:
        body.remove(elem)

    # Remove trailing empty paragraphs
    while len(doc.paragraphs) > 0:
        last_para = doc.paragraphs[-1]

        if last_para.text.strip():
            break

        p = last_para._element
        p.getparent().remove(p)

    doc.save(output_file)

    print(f"✓ Output saved as: {output_file}")


def main():

    print("\nDOCX Section Remover")
    print("--------------------\n")

    if len(sys.argv) >= 4:
        input_file = sys.argv[1]
        section_title = sys.argv[2]
        output_file = sys.argv[3]
    else:
        input_file = input("Input DOCX file: ").strip()
        section_title = input(
            "Section title to remove from (inclusive): "
        ).strip()
        output_file = input(
            "Output DOCX file: "
        ).strip()

    if not os.path.exists(input_file):
        print(f"\nERROR: File not found:\n{input_file}")
        return

    try:
        delete_from_heading_to_end(
            input_file,
            section_title,
            output_file
        )

        print("\nDone.")

    except Exception as e:
        print(f"\nERROR: {e}")


if __name__ == "__main__":
    main()
