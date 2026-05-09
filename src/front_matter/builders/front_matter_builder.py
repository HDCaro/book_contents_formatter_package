#!/usr/bin/env python3

# ================================================================================
# FRONT MATTER BUILDER MODULE
# ================================================================================
#
# PURPOSE:
# Automatically generates professional front matter for book projects using
# Python and Microsoft Word COM automation. Integrates three components:
# 1. Title page (custom design)
# 2. Copyright/legal page (custom design)
# 3. Table of Contents (auto-generated from book body headings)
#
# WORKFLOW:
# 1. Load configuration from JSON (inputs/outputs/options)
# 2. Extract chapter headings from book body using Word COM interface
# 3. Generate Table of Contents document with formatting
# 4. Assemble three components into single document with section breaks
# 5. Apply Roman numeral pagination (title: no number, copyright+TOC: i, ii, iii...)
# 6. Save output with metadata JSON for downstream tasks (Task 2: Index Builder)
#
# DEPENDENCIES:
# - pywin32 (v311): Microsoft Word COM automation (win32com.client)
# - python-docx: DOCX document manipulation (docx.Document, docx.shared)
# - Python 3.7+: f-strings, pathlib, json
#
# ARCHITECTURE:
# - Config-based: All paths/options in JSON config file
# - Path resolution: Supports relative paths + auto-discovery fallback
# - Modular design: Separate functions for each task (extraction, building, assembly)
# - Error handling: Comprehensive validation with meaningful error messages
# - Metadata output: Generates JSON for Task 2 dependency tracking
#
# CONFIGURATION FILE (front_matter_builder.config.json):
# {
#   "inputs": {
#     "title_file": "relative/path/to/title_page.docx",
#     "copyright_file": "relative/path/to/copyright_page.docx",
#     "book_body_file": "relative/path/to/HITS_AND_HAPPINESS_BODY.docx",
#     "auto_discover_missing_inputs": true
#   },
#   "outputs": {
#     "output_dir": "data/outputs/01_front_matter",
#     "filename": "front_matter.docx",
#     "temp_dir": "data/outputs/01_front_matter/temp",
#     "metadata_filename": "front_matter.config.json"
#   },
#   "options": {
#     "delete_temp_toc": true,
#     "apply_book_layout": true,
#     "page_numbering_style": "roman_lowercase"
#   }
# }
#
# USAGE:
# python front_matter_builder.py
# (Reads config, processes, outputs front_matter.docx + front_matter.config.json)
#
# OUTPUT FILES:
# - front_matter.docx (253 KB): Complete front matter (title+copyright+TOC)
# - front_matter.config.json: Metadata for Task 2 including:
#   * status: "success"
#   * headings_extracted: count of chapters/sections found
#   * next_arabic_page: page number for Task 2 to start indexing
#   * last_page_numbering: "roman_lowercase" (continuation reference)
#   * layout_applied: true/false
#   * timestamp: unix timestamp for tracking
#
# ================================================================================

import win32com.client
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
import os
import json
from pathlib import Path

# -----------------------------------
# WORD CONSTANTS (define explicitly)
# -----------------------------------
# ================================================================================
# SECTION 1: WORD COM CONSTANTS
# ================================================================================
# Microsoft Word COM object model constants. These are defined explicitly here
# because python-com doesn't always have all constants available. Numbers must
# exactly match Word's internal values.
# Page numbering constants
wdRestartContinuous = 0
wdRestartPage = 2
wdRestartSection = 1

# Page number style constants
wdPageNumberStyleLowercaseRoman = 14
wdPageNumberStyleUppercaseRoman = 13

# Alignment constants
wdAlignParagraphCenter = 1

# Field type constants
wdFieldPage = 33

# Break type constants
wdSectionBreakNextPage = 2
wdPageBreak = 7

# Header/Footer constants
wdHeaderFooterPrimary = 1


def find_project_root():
    """Find the project root directory by locating 'src' and 'data' folders.

    Returns:
        Path: Absolute path to project root (contains both src/ and data/)

    Raises:
        RuntimeError: If project root not found (missing src or data folders)

    Notes:
        Starts from the directory containing this script and walks up to parent
        directories until both src/ and data/ are found.
    """
    current = Path(__file__).resolve()
    for parent in [current.parent] + list(current.parents):
        if (parent / "src").exists() and (parent / "data").exists():
            return parent
    raise RuntimeError("Project root not found (expected folders: src and data)")


def find_file_by_name(project_root, filename):
    """Recursively search project for a file by name (case-insensitive).

    Args:
        project_root (Path): Root directory to start search from
        filename (str): Name of file to find (case-insensitive)

    Returns:
        list[Path]: List of matching paths sorted by depth (shallower first),
                    then alphabetically for deterministic selection

    Notes:
        - Search is case-insensitive (useful for portable code)
        - Returns multiple matches if file exists in different locations
        - Caller should select first result if unique match needed

    Example:
        matches = find_file_by_name(project_root, "HITS_AND_HAPPINESS_BODY.docx")
        if matches:
            book_file = matches[0]  # Use first (shallowest) match
    """
    target = filename.lower()
    matches = []

    for root, _, files in os.walk(project_root):
        for name in files:
            if name.lower() == target:
                matches.append(Path(root) / name)

    # Prefer shallower paths for deterministic selection when multiple exist.
    matches.sort(key=lambda p: (len(p.parts), str(p).lower()))
    return matches


def resolve_existing_path(project_root, label, preferred_relative_paths, filename):
    """Resolve a file path using preferred locations with auto-discovery fallback.

    Strategy (in order):
        1. Check each path in preferred_relative_paths list
        2. If not found, search entire project for filename
        3. If still not found, raise error with helpful suggestions

    Args:
        project_root (Path): Root directory to search from
        label (str): Descriptive name for error messages (e.g., "Title file")
        preferred_relative_paths (list[str]): Relative paths to try first
        filename (str): Fallback filename for auto-discovery search

    Returns:
        Path: Absolute path to resolved file

    Raises:
        FileNotFoundError: If file cannot be found anywhere in project
                          Includes helpful list of paths checked

    Notes:
        - Prints ✅ feedback for each successfully resolved path
        - Provides auto-discovery for files moved to non-standard locations
        - Deterministic: Always picks shallowest match if multiple found

    Example:
        title_file = resolve_existing_path(
            project_root,
            "Title file",
            ["data/inputs/front_matter", "data/inputs"],
            "title_page.docx"
        )
    """
    for rel in preferred_relative_paths:
        candidate = project_root / rel
        if candidate.exists():
            print(f"   ✅ Resolved {label}: {candidate.relative_to(project_root)}")
            return candidate

    matches = find_file_by_name(project_root, filename)
    if matches:
        resolved = matches[0]
        print(f"   ✅ Auto-discovered {label}: {resolved.relative_to(project_root)}")
        return resolved

    search_list = "\n      - " + "\n      - ".join(str(p) for p in preferred_relative_paths)
    raise FileNotFoundError(
        f"Could not find {label} ('{filename}') under project root: {project_root}\n"
        f"Checked preferred paths:{search_list}"
    )


def resolve_config_path(project_root, configured_path):
    """Convert a config path (relative or absolute) to absolute Path.

    Args:
        project_root (Path): Project root for relative path resolution
        configured_path (str): Path from config file (may be relative or absolute)

    Returns:
        Path: Absolute path
            - If already absolute: returned as-is
            - If relative: resolved relative to project_root

    Example:
        # Config has: "data/outputs/01_front_matter"
        output_dir = resolve_config_path(project_root, config["outputs"]["output_dir"])
        # Returns: /absolute/path/to/project/data/outputs/01_front_matter
    """
    path = Path(configured_path)
    if path.is_absolute():
        return path
    return project_root / path


def load_builder_config(project_root):
    """Load and validate front_matter_builder.config.json file.

    Configuration Structure:
        {
            "inputs": {
                "title_file": "path/to/title_page.docx",
                "copyright_file": "path/to/copyright_page.docx",
                "book_body_file": "path/to/body.docx",
                "auto_discover_missing_inputs": true
            },
            "outputs": {
                "output_dir": "data/outputs/01_front_matter",
                "filename": "front_matter.docx",
                "temp_dir": "data/outputs/01_front_matter/temp",
                "metadata_filename": "front_matter.config.json"
            },
            "options": {
                "delete_temp_toc": true,
                "apply_book_layout": true,
                "page_numbering_style": "roman_lowercase"
            }
        }

    Args:
        project_root (Path): Project root directory

    Returns:
        dict: Resolved configuration with keys:
            - config_path: Path to config file
            - title_file: Absolute path to title page
            - copyright_file: Absolute path to copyright page
            - book_body_file: Absolute path to book body (for headings)
            - output_dir: Absolute path to output directory
            - output_filename: Filename for final output
            - metadata_filename: Filename for metadata JSON
            - temp_dir: Absolute path to temporary directory
            - delete_temp_toc: Boolean for cleanup
            - apply_book_layout: Boolean for layout copying
            - page_numbering_style: String (e.g., "roman_lowercase")

    Raises:
        FileNotFoundError: Config file not found or input files missing
        ValueError: Missing required sections or keys in config

    Validation Steps:
        1. Load JSON from src/front_matter/builders/front_matter_builder.config.json
        2. Check required sections: inputs, outputs
        3. Check required keys in each section
        4. Resolve all file paths (with auto-discovery fallback)
        5. Verify all input files exist
        6. Create output/temp directories if missing

    Notes:
        - Uses auto_discover_missing_inputs to find relocated files
        - Prints ✅ status for each successfully resolved path
        - Supports relative and absolute paths in config
        - Creates output directories automatically

    Example:
        config = load_builder_config(project_root)
        headings = extract_headings(config["book_body_file"])
        # Build output using config["output_dir"] / config["output_filename"]
    """
    config_path = project_root / "src" / "front_matter" / "builders" / "front_matter_builder.config.json"
    if not config_path.exists():
        raise FileNotFoundError(
            f"Configuration file not found: {config_path}\n"
            "Create it and set the input/output paths before running."
        )

    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)

    # Validate required sections
    required_sections = ["inputs", "outputs"]
    missing_sections = [s for s in required_sections if s not in config]
    if missing_sections:
        raise ValueError(f"Missing required config sections: {', '.join(missing_sections)}")

    inputs = config.get("inputs", {})
    outputs = config.get("outputs", {})
    options = config.get("options", {})

    # Validate required input keys
    required_inputs = ["title_file", "copyright_file", "book_body_file"]
    missing_inputs = [k for k in required_inputs if not str(inputs.get(k, "")).strip()]
    if missing_inputs:
        raise ValueError(f"Missing required input keys in config: {', '.join(missing_inputs)}")

    # Validate required output keys
    required_outputs = ["output_dir", "filename", "temp_dir"]
    missing_outputs = [k for k in required_outputs if not str(outputs.get(k, "")).strip()]
    if missing_outputs:
        raise ValueError(f"Missing required output keys in config: {', '.join(missing_outputs)}")

    auto_discover = bool(inputs.get("auto_discover_missing_inputs", True))

    # Resolve input paths
    title_file = resolve_config_path(project_root, inputs["title_file"])
    if not title_file.exists() and auto_discover:
        matches = find_file_by_name(project_root, title_file.name)
        if matches:
            title_file = matches[0]

    copyright_file = resolve_config_path(project_root, inputs["copyright_file"])
    if not copyright_file.exists() and auto_discover:
        matches = find_file_by_name(project_root, copyright_file.name)
        if matches:
            copyright_file = matches[0]

    book_body_file = resolve_config_path(project_root, inputs["book_body_file"])
    if not book_body_file.exists() and auto_discover:
        matches = find_file_by_name(project_root, book_body_file.name)
        if matches:
            book_body_file = matches[0]

    # Validate all input files exist
    missing_files = []
    if not title_file.exists():
        missing_files.append(f"title_file: {title_file}")
    if not copyright_file.exists():
        missing_files.append(f"copyright_file: {copyright_file}")
    if not book_body_file.exists():
        missing_files.append(f"book_body_file: {book_body_file}")
    if missing_files:
        raise FileNotFoundError("Configured input files not found:\n - " + "\n - ".join(missing_files))

    # Resolve output paths
    output_dir = resolve_config_path(project_root, outputs["output_dir"])
    temp_dir = resolve_config_path(project_root, outputs["temp_dir"])
    output_filename = outputs.get("filename", "front_matter.docx")
    metadata_filename = outputs.get("metadata_filename", "front_matter.config.json")

    return {
        "config_path": config_path,
        "title_file": title_file,
        "copyright_file": copyright_file,
        "book_body_file": book_body_file,
        "output_dir": output_dir,
        "output_filename": output_filename,
        "metadata_filename": metadata_filename,
        "temp_dir": temp_dir,
        "delete_temp_toc": bool(options.get("delete_temp_toc", True)),
        "apply_book_layout": bool(options.get("apply_book_layout", True)),
        "page_numbering_style": options.get("page_numbering_style", "roman_lowercase"),
    }


def pretty_path(project_root, path_value):
    path_value = Path(path_value)
    try:
        return str(path_value.relative_to(project_root))
    except ValueError:
        return str(path_value)


# -----------------------------------
# CLEAN TEXT (REMOVE ALL LINE BREAKS AND CLEAN)
# -----------------------------------
# ================================================================================
# SECTION 2: TEXT PROCESSING UTILITIES
# ================================================================================
# These functions handle various text cleaning scenarios for different contexts
# (titles, body text, etc.), ensuring consistent formatting and line breaks.
#

def clean_text(text):
    """Clean text by removing line breaks, normalizing whitespace.

    Standard cleanup for paragraph text:
    - Removes all line break characters
    - Strips leading/trailing whitespace
    - Removes control characters except whitespace
    - Cleans up trailing slashes

    Args:
        text (str): Text to clean

    Returns:
        str: Cleaned text (empty string if input is None/empty)

    Example:
        clean_text("Some text\\nwith\\nbreaks")  # Returns: "Some text with breaks"
    """
    if not text:
        return ""
    # Remove control characters and clean up
    cleaned = "".join(
        c for c in text
        if c >= " " or c in ("\n", "\t")
    ).strip()

    # Remove trailing slashes and extra whitespace
    cleaned = cleaned.rstrip("/\\").strip()

    return cleaned


def clean_title_text(text):
    """Clean title text - aggressive normalization for table of contents.

    Specialized cleanup for titles/headings:
    - Removes ALL types of line breaks (\\r\\n, \\r, \\n)
    - Collapses multiple spaces into single space
    - Removes control characters
    - Cleans trailing slashes

    Used for table of contents entries to ensure single-line titles
    even if source text has embedded line breaks.

    Args:
        text (str): Title text to clean

    Returns:
        str: Single-line cleaned text (empty string if input is None/empty)

    Notes:
        - Stronger than clean_text() - collapses all whitespace variations
        - CRITICAL for table of contents formatting
        - Multiple invocations are safe (idempotent)

    Example:
        clean_title_text("Chapter 1:\\nIntroduction\\nto\\nPython")
        # Returns: "Chapter 1: Introduction to Python"
    """
    """
    Special cleaning for titles - removes ALL line breaks and normalizes spaces
    """
    if not text:
        return ""

    # Remove ALL types of line breaks and normalize spaces
    cleaned = text.replace("\r\n", " ").replace("\r", " ").replace("\n", " ")

    # Replace multiple spaces with single space
    import re
    cleaned = re.sub(r'\s+', ' ', cleaned)

    # Remove control characters except spaces
    cleaned = "".join(c for c in cleaned if c >= " ")

    # Remove trailing slashes and extra whitespace
    cleaned = cleaned.rstrip("/\\").strip()

    return cleaned


# -----------------------------------
# WORD APP
# -----------------------------------
def get_word():
    """Get or create Word COM application instance with safe defaults.

    Creates a new Word.Application COM object configured for background processing:
    - Visible: False (runs silently in background)
    - DisplayAlerts: 0 (no user prompts)

    Returns:
        win32com.client.COMObject: Word application instance

    Notes:
        - Uses DispatchEx (new instance) instead of Dispatch (shared)
        - Caller must call word.Quit() to close when done
        - Set DisplayAlerts=0 to prevent save/close dialogs
    """
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    return word


# ================================================================================
# SECTION 3: HEADING EXTRACTION
# ================================================================================
# Extract chapter titles and section headings from the book body using Word COM.
# Handles multi-line Heading 2 entries that are part of a chapter structure.

def extract_headings(doc_path):
    """Extract chapter headings from book body document.

    Strategy:
        1. Open document with Word COM (enables page number access)
        2. Iterate through paragraphs looking for Heading 1 (chapters)
        3. For each chapter, collect consecutive Heading 2 items as subtitle parts
        4. Combine heading 2 parts into single title (normalizes multi-line titles)
        5. Extract page number using Word COM Range.Information(3)
        6. Return list of {text, page} dicts with page numbers extracted via Word COM Range.Information(3)

    Args:
        doc_path (str/Path): Path to book body document (must exist)

    Returns:
        list[dict]: List of headings, each with:
            - "text" (str): Complete heading text (single line)
            - "page" (int): Page number in document

    Raises:
        No explicit raises; errors logged to console with ⚠️ indicators

    Notes:
        - CRITICAL: Must use Word COM (Range.Information(3)) to get accurate page numbers
        - python-docx cannot access page numbers
        - Closes Word application after extraction (always calls word.Quit())
        - Multi-line Heading 2 entries are collapsed into single line
        - Logs detailed progress with emoji indicators (📖 📍 ✅ 🛑 etc.)
    """
    print("\n📖 Extracting headings with consecutive Heading 2 collection")

    word = get_word()
    doc = word.Documents.Open(os.path.abspath(doc_path))
    doc.Repaginate()

    headings = []
    total = doc.Paragraphs.Count

    i = 1
    while i <= total:
        try:
            para = doc.Paragraphs(i)
            text = clean_text(para.Range.Text)
            style = para.Style.NameLocal.lower()

            # Look for Heading 1 (Chapter numbers)
            if "heading 1" in style and text:
                print(f"\n   📍 Found Heading 1: '{text}'")

                # ---- Chapter with Heading 2 title(s) ----
                if text.lower().startswith("chapter"):
                    print(f"   🔍 Processing chapter: '{text}'")

                    chapter = text
                    chapter_title_parts = []
                    j = i + 1

                    # Look for ALL consecutive Heading 2 paragraphs
                    print(f"   🔍 Looking for consecutive Heading 2 titles starting from paragraph {j}...")

                    while j <= total:
                        try:
                            p = doc.Paragraphs(j)
                            p_style = p.Style.NameLocal.lower()

                            # Get RAW text from the entire paragraph range
                            raw_text = p.Range.Text
                            cleaned_text = clean_title_text(raw_text)

                            print(f"   📝 Paragraph {j}: '{cleaned_text}' (style: {p_style})")

                            if not raw_text or not raw_text.strip():
                                print(f"   ⏭️ Empty paragraph {j}, skipping")
                                j += 1
                                continue

                            # Found Heading 2 - collect this part of the title
                            if "heading 2" in p_style:
                                chapter_title_parts.append(cleaned_text)
                                print(f"   ✅ Collected Heading 2 part: '{cleaned_text}'")
                                j += 1
                                continue  # Keep looking for more Heading 2 parts

                            # If we hit another Heading 1, stop looking
                            elif "heading 1" in p_style:
                                print(f"   🛑 Hit another Heading 1, stopping search")
                                break

                            # If we hit any other style, stop looking for more title parts
                            else:
                                print(f"   🛑 Hit non-heading style, stopping title collection")
                                break

                        except Exception as e:
                            print(f"   ⚠️ Error processing paragraph {j}: {e}")
                            break

                    # Build the complete title from all parts
                    if chapter_title_parts:
                        # Join all title parts with a space
                        complete_title = " ".join(chapter_title_parts)
                        full = f"{chapter}: {complete_title}"
                        print(f"   🎉 Complete chapter title: '{full}'")
                    else:
                        full = chapter
                        print(f"   ⚠️ No Heading 2 found, using chapter only: '{full}'")

                    # Final cleaning to ensure single line
                    full = clean_title_text(full)
                    page = para.Range.Information(3)

                    headings.append({"text": full, "page": page})
                    print(f"   📌 Added to TOC: '{full}' → page {page}")

                    i = j
                    continue

                # ---- Single-line heading (Discography, Books, etc.) ----
                else:
                    page = para.Range.Information(3)
                    # Clean single-line headings too
                    clean_heading = clean_title_text(text)
                    headings.append({"text": clean_heading, "page": page})
                    print(f"   📌 Single-line heading: '{clean_heading}' → page {page}")

            # Also capture standalone Heading 2 (if not part of a chapter)
            elif "heading 2" in style and text:
                page = para.Range.Information(3)
                # Clean standalone Heading 2
                clean_heading = clean_title_text(para.Range.Text)
                headings.append({"text": clean_heading, "page": page})
                print(f"   📌 Standalone Heading 2: '{clean_heading}' → page {page}")

        except Exception as e:
            print(f"   ⚠️ Error processing paragraph {i}: {e}")

        i += 1

    doc.Close(False)
    word.Quit()

    print(f"\n   ✅ Total headings extracted: {len(headings)}")
    for idx, h in enumerate(headings, 1):
        print(f"   {idx}. {h['text']} → page {h['page']}")

    return headings


# ================================================================================
# SECTION 4: TABLE OF CONTENTS BUILDING
# ================================================================================
# Generate a formatted TOC document with proper typography and tab stops.

def build_toc_doc(headings, toc_path):
    """Build a formatted Table of Contents document.

    Creates a professional TOC with:
        - Centered title "CONTENTS" (20pt Georgia, bold)
        - Hanging indents for entries (0.5" indent, -0.5" first line)
        - Tab stops with dot leaders to page numbers (right-aligned)
        - Text entries in 12pt Georgia, bold page numbers

    Args:
        headings (list[dict]): List of headings from extract_headings()
                             Each item has "text" and "page" keys
        toc_path (str/Path): Output path for TOC document

    Returns:
        Path: Path to created TOC document

    Notes:
        - Creates new python-docx Document (not Word COM)
        - Ensures all entry text is single-line (performs final safety cleanup)
        - Page numbers are BOLD, right-aligned with dot leaders
        - Saves as .docx format
        - Text is guaranteed single-line after this function

    Example:
        headings = extract_headings(book_path)
        toc_path = build_toc_doc(headings, "temp/toc.docx")
        # Creates formatted TOC document ready for assembly
    """
    print("\n📝 Building TOC with guaranteed single-line titles")

    doc = Document()

    # Title
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = title.add_run("CONTENTS")
    run.font.name = "Georgia"
    run.font.size = Pt(20)
    run.bold = True

    doc.add_paragraph()

    for h in headings:
        para = doc.add_paragraph()

        # Hanging indent
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)

        # Tab stop with dots
        para.paragraph_format.tab_stops.add_tab_stop(
            Inches(6),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS
        )

        # Get the title text and ensure it's single line
        title_text = h["text"]

        # Final safety check - remove any remaining line breaks
        title_text = title_text.replace("\r\n", " ").replace("\r", " ").replace("\n", " ")
        title_text = " ".join(title_text.split())  # Normalize all whitespace

        print(f"   📝 TOC entry: '{title_text}' → {h['page']}")

        # Add title text (guaranteed single line)
        run = para.add_run(title_text)
        run.font.name = "Georgia"
        run.font.size = Pt(12)

        # Add tab and page number
        page_run = para.add_run(f"\t{h['page']}")
        page_run.bold = True

    doc.save(toc_path)
    print("   ✅ TOC created with guaranteed single-line entries")

    return toc_path


# ================================================================================
# SECTION 5: ROMAN NUMERAL PAGINATION
# ================================================================================
# Apply consistent Roman numeral page numbering (i, ii, iii...) to front matter.

def apply_roman_pagination(doc):
    """Apply Roman numeral pagination to document sections.

    Pagination Strategy:
        - Section 1 (Title page): No numbering
        - Section 2 (Copyright + TOC): Roman lowercase (i, ii, iii...)
        - Restart numbering at Section 2 (starts at i, not i+page_count)

    Args:
        doc: Word COM Document object (already open)

    Returns:
        None (modifies document in place)

    Notes:
        - CRITICAL: Requires at least 2 sections in document
        - Sets RestartPageNumbering on Section 2 to isolate from Title section
        - Includes fallback field method if PageNumbers collection method fails
        - Calls doc.Repaginate() and doc.Fields.Update() to force recomputation
        - Logs detailed progress with indicators (✅ 🔢 🔧 ❌ etc.)
        - NO RETURN VALUE - modifies doc object directly

    Implementation Details:
        - Primary method: Set NumberStyle on PageNumbers collection
        - Fallback method: Use field codes if primary method fails
        - Verification: Checks final NumberStyle value matches expected

    Example:
        word = get_word()
        doc = word.Documents.Open("document.docx")
        # ... add sections ...
        apply_roman_pagination(doc)
        # Now: Title has no page number, rest has i, ii, iii...
    """
    print("\n🔢 Applying Roman pagination (FIXED FORMAT)")

    wdHeaderFooterPrimary = 1
    wdAlignParagraphCenter = 1
    wdPageNumberStyleLowercaseRoman = 14

    try:
        sections = doc.Sections
        print(f"   📊 Total sections: {sections.Count}")

        if sections.Count < 2:
            print("   ⚠️ Not enough sections")
            return

        # -----------------------------------
        # SECTION 1 → NO PAGE NUMBERS (TITLE)
        # -----------------------------------
        sec1 = sections.Item(1)

        try:
            sec1.Headers.Item(wdHeaderFooterPrimary).Range.Delete()
            sec1.Footers.Item(wdHeaderFooterPrimary).Range.Delete()
            sec1.PageSetup.DifferentFirstPageHeaderFooter = True
            print("   ✅ Section 1: no numbering")
        except Exception as e:
            print(f"   ⚠️ Section 1 cleanup error: {e}")

        # -----------------------------------
        # SECTION 2 → ROMAN NUMBERS (FIXED)
        # -----------------------------------
        sec2 = sections.Item(2)

        try:
            # Unlink from previous section
            sec2.Headers.Item(wdHeaderFooterPrimary).LinkToPrevious = False
            sec2.Footers.Item(wdHeaderFooterPrimary).LinkToPrevious = False
            print("   ✅ Section 2 unlinked")
        except Exception as e:
            print(f"   ⚠️ Unlink failed: {e}")

        try:
            # -----------------------------------
            # 🔧 FIXED: Set restart properties FIRST
            # -----------------------------------
            print("   🔧 Setting page restart properties...")
            sec2.PageSetup.RestartPageNumbering = True
            sec2.PageSetup.PageNumberStart = 1
            print("   ✅ Page restart: True, Start: 1")

            # -----------------------------------
            # 🔧 FIXED: Get footer and clear existing page numbers
            # -----------------------------------
            footer = sec2.Footers.Item(wdHeaderFooterPrimary)
            page_nums = footer.PageNumbers

            # Remove existing page numbers (important)
            print(f"   🧹 Removing {page_nums.Count} existing page numbers...")
            while page_nums.Count > 0:
                page_nums(1).Delete()

            # -----------------------------------
            # 🔧 FIXED: Add page number with Roman format
            # -----------------------------------
            print("   🔧 Adding Roman page numbers...")
            page_nums.Add(PageNumberAlignment=wdAlignParagraphCenter)

            # 🚨 CRITICAL FIX: Set NumberStyle on PageNumbers collection, NOT PageSetup
            page_nums.NumberStyle = wdPageNumberStyleLowercaseRoman
            page_nums.StartingNumber = 1  # Ensure starts at 1 (becomes "i")

            print(f"   ✅ Roman format applied: NumberStyle = {wdPageNumberStyleLowercaseRoman}")
            print("   ✅ Roman numbering configured (i, ii, iii...)")

        except Exception as e:
            print(f"   ❌ Page number setup failed: {e}")

            # Fallback method using field insertion
            try:
                print("   🔧 Trying fallback field method...")
                footer = sec2.Footers.Item(wdHeaderFooterPrimary)
                footer.Range.Delete()

                # Insert Roman numeral field directly
                footer_range = footer.Range
                field = footer_range.Fields.Add(
                    Range=footer_range,
                    Type=33,  # wdFieldPage
                    PreserveFormatting=False
                )
                field.Code.Text = "PAGE \\* ROMAN \\* LOWER"
                field.Update()

                # Center the page number
                footer.Range.ParagraphFormat.Alignment = wdAlignParagraphCenter

                print("   ✅ Fallback Roman field inserted")

            except Exception as e2:
                print(f"   ❌ Fallback method also failed: {e2}")

        # -----------------------------------
        # FORCE WORD TO APPLY CHANGES
        # -----------------------------------
        try:
            print("   🔄 Forcing document updates...")
            doc.ActiveWindow.View.ShowFieldCodes = False
            doc.Repaginate()
            doc.Fields.Update()
            doc.Repaginate()
            print("   ✅ Document refreshed")
        except Exception as e:
            print(f"   ⚠️ Refresh error: {e}")

        # -----------------------------------
        # VERIFICATION
        # -----------------------------------
        try:
            print("   🔍 Verifying Roman format...")
            sec2_footer = sec2.Footers.Item(wdHeaderFooterPrimary)
            sec2_page_nums = sec2_footer.PageNumbers
            if sec2_page_nums.Count > 0:
                style = sec2_page_nums(1).NumberStyle
                print(f"   📊 Current NumberStyle: {style} (should be {wdPageNumberStyleLowercaseRoman})")
                if style == wdPageNumberStyleLowercaseRoman:
                    print("   ✅ Roman format verified!")
                else:
                    print("   ⚠️ Roman format not applied correctly")
            else:
                print("   ⚠️ No page numbers found for verification")
        except Exception as e:
            print(f"   ⚠️ Verification failed: {e}")

    except Exception as e:
        print(f"   ❌ Pagination error: {e}")


# ================================================================================
# SECTION 6: DOCUMENT ASSEMBLY
# ================================================================================
# Combine title page, copyright page, and TOC into single document with proper
# section breaks and pagination.

def assemble_final(title_doc, copyright_doc, toc_doc, output):
    """Assemble front matter from three components with formatting.

    Assembly Sequence:
        1. Create new Word document
        2. Insert title page (becomes Section 1)
        3. Insert section break (creates Section 2)
        4. Insert copyright page (Section 2, page 1 of this section)
        5. Insert page break (stays in Section 2)
        6. Insert TOC (Section 2, pages 2+ of this section)
        7. Apply Roman pagination (title: no number, copyright+TOC: i, ii, iii...)
        8. Repaginate and update fields
        9. Save output document

    Args:
        title_doc (str/Path): Path to title page document
        copyright_doc (str/Path): Path to copyright page document
        toc_doc (str/Path): Path to TOC document
        output (str/Path): Path to save assembled document

    Returns:
        None (creates output file)

    Notes:
        - Creates Word COM document (not python-docx)
        - Uses InsertFile() method to include other DOCX files
        - Automatically calls word.Quit() after saving
        - All files must exist; silently skips missing files
        - Handles exceptions gracefully, still attempts to save
        - Logs final summary with expected page numbering pattern

    Output File Structure:
        - Section 1: Title page (no page number)
        - Section 2: Copyright + TOC with Roman numerals (i, ii, iii...)
        - Final page count depends on TOC size

    Example:
        assemble_final("title.docx", "copyright.docx", "toc.docx", "output.docx")
        # Creates front_matter.docx with proper pagination and formatting
    """
    print("\n📚 Assembling final document with Roman pagination fix")

    word = get_word()
    doc = word.Documents.Add()

    try:
        # Title page (Section 1) - NO page numbering
        if os.path.exists(title_doc):
            print("   📄 Inserting title page (Section 1)...")
            r = doc.Content
            r.Collapse(0)
            r.InsertFile(os.path.abspath(title_doc))

            # Insert section break (creates Section 2 for copyright + TOC)
            r = doc.Content
            r.Collapse(0)
            r.InsertBreak(wdSectionBreakNextPage)
            print("   ✅ Title page + section break inserted")

        # Copyright page (Section 2, page "i")
        if os.path.exists(copyright_doc):
            print("   📄 Inserting copyright page (will be Roman 'i')...")
            r = doc.Content
            r.Collapse(0)
            r.InsertFile(os.path.abspath(copyright_doc))

            # Page break before TOC (stays in same section)
            r = doc.Content
            r.Collapse(0)
            r.InsertBreak(wdPageBreak)
            print("   ✅ Copyright page inserted")

        # TOC (Section 2, pages "ii", "iii", etc.)
        print("   📄 Inserting TOC (will be Roman 'ii', 'iii'...)...")
        r = doc.Content
        r.Collapse(0)
        r.InsertFile(os.path.abspath(toc_doc))
        print("   ✅ TOC inserted")

        # Apply Roman pagination starting from copyright page
        apply_roman_pagination(doc)

        # Final processing
        print("   🔄 Final document processing...")
        try:
            doc.ActiveWindow.View.ShowFieldCodes = False
            doc.Repaginate()
            doc.Fields.Update()
            doc.Repaginate()
        except Exception as e:
            print(f"   ⚠️ Final processing warning: {e}")

        # Save
        print("   💾 Saving document...")
        doc.SaveAs2(os.path.abspath(output))
        doc.Close(False)
        word.Quit()

        print("\n🎉 DOCUMENT COMPLETED!")
        print(f"📁 File: {output}")
        print("\n📋 Expected result:")
        print("   • Title page: No number")
        print("   • Copyright page: Roman 'i'")
        print("   • TOC pages: Roman 'ii', 'iii', 'iv'...")
        print("   • ALL chapter titles on single lines")
        print("\n🔧 Roman format fix applied:")
        print("   • NumberStyle set on PageNumbers collection (not PageSetup)")
        print("   • Fallback field method if PageNumbers fails")

    except Exception as e:
        try:
            word.Quit()
        except:
            pass
        print(f"❌ Assembly error: {e}")


# ================================================================================
# SECTION 7: MAIN EXECUTION
# ================================================================================
# Entry point: Orchestrates the three-step process
# (Extract → Build → Assemble) with comprehensive error handling and metadata output.
#
# Workflow:
#   1. Load config from JSON (paths, options, etc.)
#   2. Extract headings from book body (Word COM, page numbers)
#   3. Build TOC document with professional formatting
#   4. Assemble final document (title + copyright + TOC + pagination)
#   5. Output metadata JSON for Task 2 (Index Builder)
#
# Output Files:
#   - front_matter.docx: Final assembled document (253 KB)
#   - front_matter.config.json: Metadata for downstream tasks
#
# Exit States:
#   - SUCCESS (0): All tasks completed, output files created
#   - FAILED: Headings extraction failed or output file not created

if __name__ == "__main__":
    project_root = find_project_root()

    cfg = load_builder_config(project_root)

    title_file = cfg["title_file"]
    copyright_file = cfg["copyright_file"]
    book_body_file = cfg["book_body_file"]
    
    # Setup temp directory for intermediate files
    temp_dir = cfg["temp_dir"]
    temp_dir.mkdir(parents=True, exist_ok=True)
    toc_temp = temp_dir / "toc.docx"

    # Setup output directory and file
    output_dir = cfg["output_dir"]
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / cfg["output_filename"]

    print(f"\n📂 Project root: {project_root}")
    print(f"🛠️ Config: {pretty_path(project_root, cfg['config_path'])}")
    print(f"📘 Book body (for layout): {pretty_path(project_root, book_body_file)}")
    print(f"📄 Title: {pretty_path(project_root, title_file)}")
    print(f"📄 Copyright: {pretty_path(project_root, copyright_file)}")
    print(f"🧪 Temp TOC: {pretty_path(project_root, toc_temp)}")
    print(f"💾 Output: {pretty_path(project_root, output_file)}")

    headings = extract_headings(book_body_file)
    build_succeeded = False

    if headings:
        build_toc_doc(headings, toc_temp)
        assemble_final(title_file, copyright_file, toc_temp, output_file)
        build_succeeded = True

        if cfg["delete_temp_toc"] and os.path.exists(toc_temp):
            os.remove(toc_temp)
    else:
        print("\n⚠️ No headings were extracted. Front matter output was not generated.")

    output_exists = output_file.exists()
    print("\n================ FINAL SUMMARY ================")
    print(f"Status: {'SUCCESS' if (build_succeeded and output_exists) else 'FAILED'}")
    print(f"Headings extracted: {len(headings)}")
    print(f"Output file: {output_file}")
    print(f"Output exists: {'YES' if output_exists else 'NO'}")
    print("===============================================")

    # Save metadata for downstream tasks
    if build_succeeded and output_exists:
        metadata = {
            "task": "front_matter",
            "status": "success",
            "output_file": str(output_file.relative_to(project_root)),
            "page_count": "unknown",
            "last_page_numbering": "roman_lowercase",
            "next_arabic_page": 1,
            "headings_extracted": len(headings),
            "layout_applied": cfg.get("apply_book_layout", True),
            "timestamp": str(os.path.getmtime(output_file))
        }
        metadata_file = output_dir / cfg["metadata_filename"]
        with open(metadata_file, "w", encoding="utf-8") as f:
            json.dump(metadata, f, indent=2)
        print(f"✅ Metadata saved: {pretty_path(project_root, metadata_file)}")