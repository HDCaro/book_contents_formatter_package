import re
import json
import sys
from pathlib import Path
from collections import defaultdict

import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.opc.exceptions import PackageNotFoundError

# 🔧 REQUIRED FOR COLUMNS
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------- CONFIG ---------------- #

DOCX_INPUT = "HITS AND HAPPINESS FINAL 2 Format MOM Discog.docx"
DOCX_OUTPUT = Path(DOCX_INPUT).with_name(
    Path(DOCX_INPUT).stem + " index.docx"
)

CACHE_FILE = "discogs_cache.json"

WORDS_PER_PAGE = 600
MIN_OCCURRENCES = 2

DISCOGS_TOKEN = ""
VERBOSE = True

# ---------------- LOGGING ---------------- #

def log(msg):
    if VERBOSE:
        print(msg)

# ---------------- CACHE ---------------- #

def load_cache():
    try:
        with open(CACHE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except:
        return {}

def save_cache(cache):
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, indent=2)

cache = load_cache()

# ---------------- TEXT ---------------- #

def extract_text(docx):
    try:
        doc = Document(docx)
        return "\n".join(p.text for p in doc.paragraphs)
    except PackageNotFoundError:
        print(f"❌ File not found: {docx}")
        sys.exit(1)

def split_pages(text):
    words = text.split()
    pages = [
        " ".join(words[i:i + WORDS_PER_PAGE])
        for i in range(0, len(words), WORDS_PER_PAGE)
    ]
    log(f"[INFO] Total pages estimated: {len(pages)}")
    return pages

# ---------------- NORMALIZATION ---------------- #

def normalize_title(text):
    if "," in text:
        parts = [p.strip() for p in text.split(",")]
        parts.reverse()
        return " ".join(parts)
    return text

def clean_query(text):
    return re.sub(r"[^\w\s]", "", text)

# ---------------- DISCOGS ---------------- #

def discogs_search(query):
    query = clean_query(normalize_title(query))

    if query in cache:
        return cache[query]

    headers = {"User-Agent": "IndexBuilder/1.0"}
    if DISCOGS_TOKEN:
        headers["Authorization"] = f"Discogs token={DISCOGS_TOKEN}"

    for q in [query, f"{query} song", f"{query} track"]:
        try:
            r = requests.get(
                "https://api.discogs.com/database/search",
                params={"q": q},
                headers=headers,
                timeout=5
            )
            results = r.json().get("results", [])
            if results:
                cache[query] = True
                return True
        except Exception as e:
            log(f"[ERROR] Discogs request failed: {e}")

    cache[query] = False
    return False

# ---------------- DETECTION ---------------- #

def extract_candidates(text):
    pattern = r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\b"
    return re.findall(pattern, text)

# ---------------- CLASSIFICATION ---------------- #

BAND_KEYWORDS = {
    "Band", "Boys", "Girls", "Group", "Orchestra",
    "Ensemble", "Quartet", "Trio", "Duo",
    "Project", "Collective", "Crew", "Gang"
}

def is_likely_band(parts):
    if any(p in BAND_KEYWORDS for p in parts):
        return True
    if parts[0] == "The":
        return True
    if len(parts) >= 3:
        return True
    return False

def is_likely_person(parts):
    if len(parts) in (2, 3):
        if not any(p in BAND_KEYWORDS for p in parts):
            if parts[0] != "The":
                return True
    return False

def invert_person(parts):
    return f"{parts[-1]}, {' '.join(parts[:-1])}"

def normalize_band_name(parts):
    if parts[0] == "The" and len(parts) > 1:
        return f"{' '.join(parts[1:])}, The"
    return " ".join(parts)

def classify_entity(candidate):
    parts = candidate.split()
    normalized = normalize_title(candidate)

    # 1. WORK (Discogs)
    if discogs_search(normalized):
        return normalized

    # 2. BAND
    if is_likely_band(parts):
        return normalize_band_name(parts)

    # 3. PERSON
    if is_likely_person(parts):
        return invert_person(parts)

    # 4. FALLBACK (no inversion)
    return candidate

def classify_entries(text):
    candidates = extract_candidates(text)
    results = []

    for c in candidates:
        parts = c.split()

        if len(parts) < 2:
            continue

        classified = classify_entity(c)
        results.append(classified)

    return results

# ---------------- INDEX ---------------- #

def build_index(pages):
    index = defaultdict(set)

    for i, page in enumerate(pages, 1):
        entries = classify_entries(page)
        for entry in entries:
            index[entry].add(i)

    return {
        k: v for k, v in index.items()
        if len(v) >= MIN_OCCURRENCES
    }

def compress(pages):
    pages = sorted(pages)
    ranges = []
    start = prev = pages[0]

    for p in pages[1:]:
        if p == prev + 1:
            prev = p
        else:
            ranges.append((start, prev))
            start = prev = p

    ranges.append((start, prev))

    return ", ".join(
        f"{a}–{b}" if a != b else str(a)
        for a, b in ranges
    )

def build_alphabetical_index(index):
    grouped = defaultdict(list)

    for entry in sorted(index):
        line = f"{entry}, {compress(list(index[entry]))}"
        letter = entry[0].upper()
        grouped[letter].append(line)

    return dict(sorted(grouped.items()))

# ---------------- DOCX ---------------- #

def set_two_columns(doc):
    section = doc.sections[0]
    sectPr = section._sectPr

    cols = sectPr.xpath('./w:cols')
    if cols:
        cols[0].set(qn('w:num'), "2")
    else:
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), "2")
        sectPr.append(cols)

def create_doc(alpha_index, out):
    doc = Document()

    # Title
    title = doc.add_paragraph()
    r = title.add_run("INDEX")
    r.bold = True
    r.font.name = "Georgia"
    r.font.size = Pt(20)
    title.alignment = 1

    set_two_columns(doc)

    for letter, entries in alpha_index.items():

        p = doc.add_paragraph()
        r = p.add_run(letter)
        r.bold = True
        r.font.name = "Georgia"
        r.font.size = Pt(14)

        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)

        for e in entries:
            ep = doc.add_paragraph(e)
            ep.paragraph_format.left_indent = Inches(0.3)
            ep.paragraph_format.first_line_indent = Inches(-0.3)
            ep.paragraph_format.space_before = Pt(0)
            ep.paragraph_format.space_after = Pt(0)

    try:
        doc.save(out)
    except PermissionError:
        print("\n❌ Cannot save file: it is currently in use.")
        print(f"👉 {out}")
        print("\nPlease close the file in Word and try again.")
        sys.exit(1)

# ---------------- MAIN ---------------- #

def main():
    print("\n=== START INDEX GENERATION ===\n")

    text = extract_text(DOCX_INPUT)
    pages = split_pages(text)

    index = build_index(pages)
    alpha_index = build_alphabetical_index(index)

    create_doc(alpha_index, DOCX_OUTPUT)

    save_cache(cache)

    print("\n=== DONE ===")

if __name__ == "__main__":
    main()