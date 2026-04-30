#!/usr/bin/env python3
"""
Table of Contents Formatter - FIXED VERSION
Properly aligns page numbers to the right margin using tab stops
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
import re

def create_formatted_contents_fixed():
    """Create a professionally formatted table of contents with proper right-aligned page numbers"""
    
    # Raw contents data extracted from the document
    contents_data = [
        ("Introduction", 1),
        ("Chapter 1", "'Pre' His Story", 5),
        ("Chapter 2", "What is an \"Arranger\"?", 17),
        ("Chapter 3", "Kid Stuff", 27),
        ("Chapter 4", "School Daze", 36),
        ("Chapter 5", "Pure Wings", 50),
        ("Chapter 6", "Berklee Bodhisattva", 58),
        ("Chapter 7", "London 1975 – You CAN Go Home Again", 72),
        ("Chapter 8", "Onward and Upward", 89),
        ("Chapter 9", "Interplanetary Bollocks", 103),
        ("Chapter 10", "Say Hello to Hollywood", 119),
        ("Chapter 11", "From the Garage, Back to London", 130),
        ("Chapter 12", "Finally – The '80s", 142),
        ("Chapter 13", "Heeeeeeeeee's Gracie", 163),
        ("Chapter 14", "The Birth of Bandzilla and Other Tails", 178),
        ("Chapter 15", "Croissants, Butterflies, and Breakout", 182),
        ("Chapter 16", "Hollywood in Notting Hill and Breaking out in Chiswick", 198),
        ("Chapter 17", "Bandzilla Devours a Star", 210),
        ("Chapter 18", "Pet Shop Boys, Holly and the Seeds of Love", 220),
        ("Chapter 19", "Tell Me Where You're Going", 238),
        ("Chapter 20", "Three Boys in a Pet Shop", 257),
        ("Chapter 21", "Been Around The World...", 273),
        ("Chapter 22", "Radio Richard", 288),
        ("Chapter 23", "Santa Rita, Pray For Me", 296),
        ("Chapter 24", "Pat Kane Meets Zilla", 301),
        ("Chapter 25", "Work, Work, Work", 308),
        ("Chapter 26", "Brother Ray", 324),
        ("Chapter 27", "Pop Music", 334),
        ("Chapter 28", "Wet Hits and Having a Ball", 347),
        ("Chapter 29", "Four Artists", 357),
        ("Chapter 30", "Work, Work, Play", 367),
        ("Chapter 31", "The Cheery '00s", 377),
        ("Chapter 32", "California, There I Came", 389),
        ("Chapter 33", "It's an Education", 400),
        ("Chapter 34", "Smiling at Tombs, and a Classical Gas", 404),
        ("Chapter 35", "Thoughts", 414)
    ]
    
    # Create new document
    doc = Document()
    
    # Set up page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("•CONTENTS•")
    title_run.font.name = 'Georgia'
    title_run.font.size = Pt(16)
    title_run.bold = True
    
    # Add some space after title
    doc.add_paragraph()
    
    # Process each entry
    for entry in contents_data:
        if len(entry) == 2:  # Introduction case
            title, page = entry
            
            # Add chapter title (Georgia 20 bold, centered)
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            title_run.font.name = 'Georgia'
            title_run.font.size = Pt(20)
            title_run.bold = True
            
            # Add page number line with proper right alignment
            page_para = doc.add_paragraph()
            page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set up tab stop for right-aligned page number
            tab_stops = page_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            
            page_run = page_para.add_run(f"\t{page}")
            page_run.font.name = 'Georgia'
            page_run.font.size = Pt(12)
            
        else:  # Chapter cases
            chapter_num, chapter_title, page = entry
            
            # Add chapter number (Georgia 12, centered)
            num_para = doc.add_paragraph()
            num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            num_run = num_para.add_run(chapter_num)
            num_run.font.name = 'Georgia'
            num_run.font.size = Pt(12)
            
            # Add chapter title (Georgia 20 bold, centered)
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(chapter_title)
            title_run.font.name = 'Georgia'
            title_run.font.size = Pt(20)
            title_run.bold = True
            
            # Add page number line with proper right alignment
            page_para = doc.add_paragraph()
            page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set up tab stop for right-aligned page number
            tab_stops = page_para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            
            page_run = page_para.add_run(f"\t{page}")
            page_run.font.name = 'Georgia'
            page_run.font.size = Pt(12)
        
        # Add spacing between entries
        doc.add_paragraph()
    
    return doc

def create_traditional_format_fixed():
    """Create traditional format with PROPER right-aligned page numbers"""
    
    doc = Document()
    
    # Set up page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("•CONTENTS•")
    title_run.font.name = 'Georgia'
    title_run.font.size = Pt(16)
    title_run.bold = True
    
    doc.add_paragraph()
    
    # Contents data
    contents_data = [
        ("Introduction", 1),
        ("Chapter 1: 'Pre' His Story", 5),
        ("Chapter 2: What is an \"Arranger\"?", 17),
        ("Chapter 3: Kid Stuff", 27),
        ("Chapter 4: School Daze", 36),
        ("Chapter 5: Pure Wings", 50),
        ("Chapter 6: Berklee Bodhisattva", 58),
        ("Chapter 7: London 1975 – You CAN Go Home Again", 72),
        ("Chapter 8: Onward and Upward", 89),
        ("Chapter 9: Interplanetary Bollocks", 103),
        ("Chapter 10: Say Hello to Hollywood", 119),
        ("Chapter 11: From the Garage, Back to London", 130),
        ("Chapter 12: Finally – The '80s", 142),
        ("Chapter 13: Heeeeeeeeee's Gracie", 163),
        ("Chapter 14: The Birth of Bandzilla and Other Tails", 178),
        ("Chapter 15: Croissants, Butterflies, and Breakout", 182),
        ("Chapter 16: Hollywood in Notting Hill and Breaking out in Chiswick", 198),
        ("Chapter 17: Bandzilla Devours a Star", 210),
        ("Chapter 18: Pet Shop Boys, Holly and the Seeds of Love", 220),
        ("Chapter 19: Tell Me Where You're Going", 238),
        ("Chapter 20: Three Boys in a Pet Shop", 257),
        ("Chapter 21: Been Around The World...", 273),
        ("Chapter 22: Radio Richard", 288),
        ("Chapter 23: Santa Rita, Pray For Me", 296),
        ("Chapter 24: Pat Kane Meets Zilla", 301),
        ("Chapter 25: Work, Work, Work", 308),
        ("Chapter 26: Brother Ray", 324),
        ("Chapter 27: Pop Music", 334),
        ("Chapter 28: Wet Hits and Having a Ball", 347),
        ("Chapter 29: Four Artists", 357),
        ("Chapter 30: Work, Work, Play", 367),
        ("Chapter 31: The Cheery '00s", 377),
        ("Chapter 32: California, There I Came", 389),
        ("Chapter 33: It's an Education", 400),
        ("Chapter 34: Smiling at Tombs, and a Classical Gas", 404),
        ("Chapter 35: Thoughts", 414)
    ]
    
    # Add each entry with PROPER right-aligned page numbers
    for title, page in contents_data:
        para = doc.add_paragraph()
        
        # Set up tab stop for right-aligned page number with dot leaders
        # Position the tab stop near the right margin
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        # Add title
        title_run = para.add_run(title)
        title_run.font.name = 'Georgia'
        title_run.font.size = Pt(12)
        
        # Add tab and page number (this will be right-aligned with dot leaders)
        page_run = para.add_run(f"\t{page}")
        page_run.font.name = 'Georgia'
        page_run.font.size = Pt(12)
        page_run.bold = True
    
    return doc

def update_advanced_indexer_alignment():
    """Update the advanced indexer to use proper right alignment"""
    
    # Read the current advanced indexer
    with open('advanced_book_indexer.py', 'r') as f:
        content = f.read()
    
    # Replace the problematic dot calculation with proper tab stops
    old_code = '''            # Calculate spacing for dots
            # This creates a right-aligned page number with dot leaders
            title_length = len(title)
            available_space = 80  # Approximate character width
            dots_needed = max(available_space - title_length - len(str(page)), 3)
            
            # Add dots
            dots_run = para.add_run("." * dots_needed)
            dots_run.font.name = 'Georgia'
            dots_run.font.size = Pt(12)
            
            # Add page number
            page_run = para.add_run(str(page))
            page_run.font.name = 'Georgia'
            page_run.font.size = Pt(12)
            page_run.bold = True'''
    
    new_code = '''            # Set up tab stop for right-aligned page number with dot leaders
            from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
            tab_stops = para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            
            # Add tab and page number (this will be right-aligned with dot leaders)
            page_run = para.add_run(f"\\t{page}")
            page_run.font.name = 'Georgia'
            page_run.font.size = Pt(12)
            page_run.bold = True'''
    
    # Replace the old code with new code
    updated_content = content.replace(old_code, new_code)
    
    # Also need to add the import at the top
    if "from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER" not in updated_content:
        # Find the existing imports and add the new one
        import_line = "from docx.enum.text import WD_ALIGN_PARAGRAPH"
        if import_line in updated_content:
            updated_content = updated_content.replace(
                import_line, 
                "from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER"
            )
    
    # Save the updated file
    with open('advanced_book_indexer_fixed.py', 'w') as f:
        f.write(updated_content)

if __name__ == "__main__":
    print("Creating FIXED formatted table of contents with proper right-aligned page numbers...")
    
    # Create the author's preferred format with FIXED alignment
    print("1. Creating author's preferred format (FIXED alignment)...")
    doc1 = create_formatted_contents_fixed()
    doc1.save("contents_author_format_FIXED.docx")
    print("   ✓ Saved as 'contents_author_format_FIXED.docx'")
    
    # Create traditional format with FIXED alignment
    print("2. Creating traditional format (FIXED alignment)...")
    doc2 = create_traditional_format_fixed()
    doc2.save("contents_traditional_format_FIXED.docx")
    print("   ✓ Saved as 'contents_traditional_format_FIXED.docx'")
    
    # Update the advanced indexer
    print("3. Creating fixed version of advanced indexer...")
    try:
        update_advanced_indexer_alignment()
        print("   ✓ Saved as 'advanced_book_indexer_fixed.py'")
    except Exception as e:
        print(f"   ⚠ Could not update indexer: {e}")
    
    print("\n🎉 FIXED versions created successfully!")
    print("\n✅ Key improvements:")
    print("- Page numbers are now properly right-aligned to the margin")
    print("- Uses Word's built-in tab stops with dot leaders")
    print("- Professional typography and spacing")
    print("- Consistent alignment across all entries")
    
    print("\n📁 Files created:")
    print("- contents_author_format_FIXED.docx (author's preferred style)")
    print("- contents_traditional_format_FIXED.docx (traditional style)")
    print("- advanced_book_indexer_fixed.py (updated script)")