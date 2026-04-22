#!/usr/bin/env python3
"""
Test Example for Book Contents Formatter
Creates a sample book with headings to test the indexer
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_book():
    """Create a sample book with proper heading styles for testing"""
    
    doc = Document()
    
    # Set up margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    
    # Title page
    title = doc.add_heading('Sample Book: A Test Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Chapter 1
    doc.add_heading('Chapter 1: Introduction', 1)
    doc.add_paragraph('This is the introduction chapter. It contains important information about the book.')
    doc.add_paragraph('Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.')
    
    # Subsections
    doc.add_heading('What You Will Learn', 2)
    doc.add_paragraph('In this book, you will discover many interesting topics.')
    
    doc.add_heading('How to Use This Book', 2)
    doc.add_paragraph('Each chapter builds upon the previous one.')
    
    # Chapter 2
    doc.add_heading('Chapter 2: Getting Started', 1)
    doc.add_paragraph('This chapter covers the basics of getting started.')
    doc.add_paragraph('We will walk through step-by-step instructions.')
    
    doc.add_heading('Prerequisites', 2)
    doc.add_paragraph('Before you begin, make sure you have the following.')
    
    doc.add_heading('Installation', 2)
    doc.add_paragraph('Follow these steps to install the required software.')
    
    doc.add_heading('Configuration', 3)
    doc.add_paragraph('Configure your environment with these settings.')
    
    # Chapter 3
    doc.add_heading('Chapter 3: Advanced Topics', 1)
    doc.add_paragraph('This chapter covers more advanced concepts.')
    
    doc.add_heading('Best Practices', 2)
    doc.add_paragraph('Follow these best practices for optimal results.')
    
    doc.add_heading('Troubleshooting', 2)
    doc.add_paragraph('Common issues and their solutions.')
    
    # Chapter 4
    doc.add_heading('Chapter 4: Conclusion', 1)
    doc.add_paragraph('In conclusion, this book has covered many important topics.')
    doc.add_paragraph('Thank you for reading!')
    
    # Save the sample book
    doc.save('sample_book_body.docx')
    print("✓ Created 'sample_book_body.docx' for testing")
    print("✓ Contains 4 chapters with proper heading styles")
    print("✓ Ready to test with advanced_book_indexer.py")

def test_indexer():
    """Test the indexer with the sample book"""
    try:
        from advanced_book_indexer import create_book_with_index
        
        print("\nTesting the indexer...")
        headings = create_book_with_index(
            'sample_book_body.docx', 
            'test_output_with_index.docx',
            font_name="Georgia",
            chapter_num_size=12,
            chapter_title_size=20,
            alignment="center"
        )
        
        print(f"\n✓ Test completed successfully!")
        print(f"✓ Generated index with {len(headings)} headings")
        print(f"✓ Output saved as 'test_output_with_index.docx'")
        
    except ImportError:
        print("❌ Could not import advanced_book_indexer")
        print("Make sure advanced_book_indexer.py is in the same folder")
    except Exception as e:
        print(f"❌ Error during testing: {e}")

if __name__ == "__main__":
    print("Book Contents Formatter - Test Example")
    print("=====================================")
    print()
    
    # Create sample book
    create_sample_book()
    
    # Test the indexer
    test_indexer()
    
    print("\nFiles created:")
    print("- sample_book_body.docx (test input)")
    print("- test_output_with_index.docx (test output)")
    print()
    print("Next steps:")
    print("1. Open test_output_with_index.docx to see the result")
    print("2. If it looks good, use your real book file")
    print("3. Customize the formatting as needed")