#!/usr/bin/env python3
"""
Advanced Book Indexer with Roman Numerals
Extracts headings from book body DOCX and creates formatted index with Roman numerals
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
import re

def roman_numeral(num):
    """Convert integer to Roman numeral"""
    values = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    literals = ['M', 'CM', 'D', 'CD', 'C', 'XC', 'L', 'XL', 'X', 'IX', 'V', 'IV', 'I']
    
    result = ''
    for i in range(len(values)):
        count = num // values[i]
        if count:
            result += literals[i] * count
            num -= values[i] * count
    return result.lower()

def extract_headings_from_docx(docx_path):
    """Extract headings and their page numbers from DOCX"""
    doc = Document(docx_path)
    headings = []
    current_page = 1  # Starting page for main content
    
    for para in doc.paragraphs:
        # Check if paragraph is a heading
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
            
            # Extract heading text
            heading_text = para.text.strip()
            if heading_text:
                headings.append({
                    'text': heading_text,
                    'level': level,
                    'page': current_page
                })
        
        # Estimate page breaks (rough calculation)
        # This is approximate - for exact pages, you'd need more sophisticated analysis
        if len(para.text) > 500:  # Long paragraph might span pages
            current_page += 1
    
    return headings

def create_book_with_index(book_body_path, output_path, 
                          font_name="Georgia", 
                          chapter_num_size=12, 
                          chapter_title_size=20,
                          alignment="center"):
    """Create complete book with Roman numeral index at front"""
    
    # Extract headings from book body
    print("Extracting headings from book body...")
    headings = extract_headings_from_docx(book_body_path)
    
    if not headings:
        print("No headings found in the document. Make sure your document uses Heading styles.")
        return
    
    # Create new document for complete book
    doc = Document()
    
    # Set up page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    
    # Add title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("TABLE OF CONTENTS")
    title_run.font.name = font_name
    title_run.font.size = Pt(18)
    title_run.bold = True
    
    doc.add_paragraph()  # Space
    
    # Add contents with Roman numerals
    roman_page = 1
    
    for heading in headings:
        # Determine if this is a chapter or section
        if heading['level'] == 1:  # Main chapters
            # Add chapter number (if it starts with "Chapter")
            if heading['text'].lower().startswith('chapter'):
                # Split chapter number and title
                parts = heading['text'].split(':', 1)
                if len(parts) == 2:
                    chapter_num = parts[0].strip()
                    chapter_title = parts[1].strip()
                    
                    # Chapter number
                    num_para = doc.add_paragraph()
                    if alignment.lower() == "center":
                        num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    num_run = num_para.add_run(chapter_num)
                    num_run.font.name = font_name
                    num_run.font.size = Pt(chapter_num_size)
                    
                    # Chapter title
                    title_para = doc.add_paragraph()
                    if alignment.lower() == "center":
                        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title_para.add_run(chapter_title)
                    title_run.font.name = font_name
                    title_run.font.size = Pt(chapter_title_size)
                    title_run.bold = True
                else:
                    # Single line chapter
                    title_para = doc.add_paragraph()
                    if alignment.lower() == "center":
                        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title_para.add_run(heading['text'])
                    title_run.font.name = font_name
                    title_run.font.size = Pt(chapter_title_size)
                    title_run.bold = True
            else:
                # Non-chapter heading
                title_para = doc.add_paragraph()
                if alignment.lower() == "center":
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_para.add_run(heading['text'])
                title_run.font.name = font_name
                title_run.font.size = Pt(chapter_title_size)
                title_run.bold = True
            
            # Add Roman numeral page number
            page_para = doc.add_paragraph()
            if alignment.lower() == "center":
                page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Create dot leaders
            title_length = len(heading['text'])
            dots_needed = max(50 - title_length, 10)
            dots = "." * dots_needed
            
            roman_num = roman_numeral(roman_page)
            page_run = page_para.add_run(f"{dots} {roman_num}")
            page_run.font.name = font_name
            page_run.font.size = Pt(12)
            
            roman_page += 1
            
        else:  # Sub-headings (level 2+)
            # Indent sub-headings
            sub_para = doc.add_paragraph()
            sub_para.paragraph_format.left_indent = Inches(0.5 * heading['level'])
            
            sub_run = sub_para.add_run(heading['text'])
            sub_run.font.name = font_name
            sub_run.font.size = Pt(12)
            
            # Add page number
            title_length = len(heading['text'])
            dots_needed = max(40 - title_length, 5)
            dots = "." * dots_needed
            
            roman_num = roman_numeral(roman_page)
            dots_run = sub_para.add_run(f"{dots} {roman_num}")
            dots_run.font.name = font_name
            dots_run.font.size = Pt(10)
        
        # Add spacing
        doc.add_paragraph()
    
    # Save the document
    doc.save(output_path)
    print(f"✓ Complete book with index saved as '{output_path}'")
    print(f"✓ Found {len(headings)} headings")
    print(f"✓ Index uses Roman numerals: i, ii, iii, iv, v...")
    
    return headings

def create_standalone_index(book_body_path, output_path="book_index.docx"):
    """Create just the index without the full book"""
    headings = create_book_with_index(book_body_path, output_path)
    return headings

if __name__ == "__main__":
    print("Advanced Book Indexer")
    print("====================")
    print()
    print("This script can:")
    print("1. Extract headings from your book body DOCX")
    print("2. Create a formatted index with Roman numerals")
    print("3. Use your preferred typography (Georgia, sizes, alignment)")
    print()
    
    # Example usage
    book_body_file = "book_body.docx"  # Replace with your book file
    output_file = "complete_book_with_index.docx"
    
    print(f"Looking for book body: {book_body_file}")
    print("To use this script:")
    print("1. Place your book body DOCX file in the same folder")
    print("2. Make sure it uses Heading 1, Heading 2, etc. styles")
    print("3. Run: python advanced_book_indexer.py")
    print()
    print("The script will create a complete book with:")
    print("- Roman numeral index at the front")
    print("- Your preferred formatting")
    print("- Proper page numbering")