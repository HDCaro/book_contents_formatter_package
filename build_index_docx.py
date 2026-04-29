"""
===============================================================================
FILE: build_index_docx.py
-------------------------------------------------------------------------------
Builds final DOCX index with:

- merge logic
- aliases + aliases_external (single AKA line)
- filtering (MIN_PAGES)
- professional alphabetical sorting
- SAFE JSON serialization
- canonical alias filtering (handles reversed names)
===============================================================================
"""

import json
import re
import unicodedata
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------- CONFIG ---------------- #

DOCX_INPUT = "HITS AND HAPPINESS FINAL 2 Format MOM Discog.docx"
RAW_JSON = "index_raw.json"
CURATED_JSON = "index_curated.json"
FILTERED_JSON = "index_filtered_out.json"

MIN_PAGES = 2

DOCX_OUTPUT = Path(DOCX_INPUT).with_name(
    Path(DOCX_INPUT).stem + "-index.docx"
)

# ---------------- SAFETY ---------------- #

def serialize_entry(v):
    return {
        "pages": sorted(list(v.get("pages", []))),
        "aliases": sorted(list(v.get("aliases", []))),
        "aliases_external": sorted(list(v.get("aliases_external", [])))
    }

# ---------------- NAME NORMALIZATION ---------------- #

def normalize_person_name(name):
    """
    Normalize identity so:
    'Silver-Lasky, Pat' == 'Pat Silver-Lasky'
    """
    name = name.lower().strip()

    name = name.replace("’", "'")

    # Last, First → First Last
    if "," in name:
        parts = [p.strip() for p in name.split(",")]
        if len(parts) == 2:
            name = f"{parts[1]} {parts[0]}"

    name = re.sub(r"\s+", " ", name)

    return name

# ---------------- CURATION ---------------- #

def apply_curation(raw, curated):
    final = deepcopy(raw)

    for v in final.values():
        v["pages"] = set(v["pages"])

    # remove
    for k, r in curated.items():
        if r.get("action") == "remove" and k in final:
            del final[k]

    # merge
    for k, r in curated.items():
        if r.get("action") == "merged":
            dest = r.get("destination")

            if dest not in final:
                final[dest] = {
                    "normalized": dest,
                    "type": "unknown",
                    "pages": set()
                }

            if k in final:
                final[dest]["pages"].update(final[k]["pages"])
                del final[k]

    # enforce normalized
    for k, r in curated.items():
        if k in final and "normalized" in r:
            final[k]["normalized"] = r["normalized"]

    return final

# ---------------- NORMALIZED INDEX ---------------- #

def build_normalized_index(final, curated):
    normalized_index = {}

    for key, v in final.items():
        norm = v["normalized"]

        normalized_index.setdefault(norm, {
            "pages": set(),
            "aliases": set(),
            "aliases_external": set()
        })

        normalized_index[norm]["pages"].update(v["pages"])

    # merged aliases
    for ck, rule in curated.items():
        if rule.get("action") == "merged":
            dest = rule.get("destination")

            if not dest or dest not in final:
                continue

            dest_norm = final[dest]["normalized"]
            normalized_index[dest_norm]["aliases"].add(ck)

    # curated overrides + enrichments
    for key, rule in curated.items():
        if key in final:
            norm = final[key]["normalized"]

            # override pages
            if "pages" in rule:
                normalized_index[norm]["pages"] = set(rule["pages"])

            for a in rule.get("aliases", []):
                normalized_index[norm]["aliases"].add(a)

            for a in rule.get("aliases_external", []):
                normalized_index[norm]["aliases_external"].add(a)

    return normalized_index

# ---------------- SORTING ---------------- #

def strip_accents(text):
    return ''.join(
        c for c in unicodedata.normalize('NFD', text)
        if unicodedata.category(c) != 'Mn'
    )

def normalize_sort_key(text):
    t = text.lower().strip()

    t = re.sub(r"^(the|a|an)\s+", "", t)
    t = t.replace("’", "'")
    t = strip_accents(t)

    t = re.sub(r"\bmc([a-z])", r"mac\1", t)
    t = re.sub(r"\b([od])['’]([a-z])", r"\1\2", t)

    return t

def get_sort_value(entry):
    return entry["name"] if entry["type"] == "alias" else entry["line"].split(",")[0]

def sort_entries(entries):
    return sorted(entries, key=lambda e: normalize_sort_key(get_sort_value(e)))

# ---------------- UTIL ---------------- #

def compress(pages):
    pages = sorted(pages)
    ranges = []
    start = prev = pages[0]

    for p in pages[1:]:
        if p == prev + 1:
            prev = p
        else:
            ranges.append((start, prev))
            start = prev = p

    ranges.append((start, prev))

    return ", ".join(
        f"{a}–{b}" if a != b else str(a)
        for a, b in ranges
    )

# ---------------- BUILD ---------------- #

def build_alpha(index):
    grouped = {}
    filtered = {}

    for name, v in index.items():

        if len(v["pages"]) < MIN_PAGES:
            filtered[name] = serialize_entry(v)
            continue

        pages = compress(v["pages"])
        letter = name[0].upper()

        canonical_norm = normalize_person_name(name)

        all_aliases = sorted(
            a for a in (set(v["aliases"]) | set(v["aliases_external"]))
            if normalize_person_name(a) != canonical_norm
        )

        entry = {
            "type": "alias" if all_aliases else "normal",
            "name": name,
            "aliases": "; ".join(all_aliases),
            "pages": pages
        }

        if not all_aliases:
            entry["line"] = f"{name}, {pages}"

        grouped.setdefault(letter, []).append(entry)

    # sort entries
    for letter in grouped:
        grouped[letter] = sort_entries(grouped[letter])

    with open(FILTERED_JSON, "w", encoding="utf-8") as f:
        json.dump(filtered, f, indent=2)

    return dict(sorted(grouped.items()))

# ---------------- DOCX ---------------- #

def set_two_columns(doc):
    section = doc.sections[0]
    sectPr = section._sectPr

    cols = sectPr.xpath('./w:cols')
    if cols:
        cols[0].set(qn('w:num'), "2")
    else:
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), "2")
        sectPr.append(cols)

def create_doc(alpha):
    doc = Document()

    title = doc.add_paragraph()
    run = title.add_run("INDEX")
    run.bold = True
    run.font.size = Pt(20)
    title.alignment = 1

    set_two_columns(doc)

    for letter, entries in alpha.items():
        p = doc.add_paragraph()

        run = p.add_run(letter)
        run.bold = True
        run.font.size = Pt(14)
        run.add_break()

        for entry in entries:

            if entry["type"] == "alias":
                p.add_run(entry["name"])
                p.add_run().add_break()

                if entry.get("aliases"):
                    p.add_run(f"(AKA: {entry['aliases']})")
                    p.add_run().add_break()

                p.add_run(entry["pages"])
                p.add_run().add_break()

            else:
                p.add_run(entry["line"])
                p.add_run().add_break()

    doc.save(DOCX_OUTPUT)

# ---------------- MAIN ---------------- #

def main():
    raw = json.load(open(RAW_JSON))
    curated = json.load(open(CURATED_JSON))

    final = apply_curation(raw, curated)
    index = build_normalized_index(final, curated)

    alpha = build_alpha(index)
    create_doc(alpha)

    print("\n✅ Index created:", DOCX_OUTPUT)

if __name__ == "__main__":
    main()